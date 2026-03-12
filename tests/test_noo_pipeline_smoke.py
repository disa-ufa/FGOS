from __future__ import annotations

from docx import Document

from shared.rubric.load import load_rubric
from worker.pipeline.noo_extract import extract_noo_from_canonical
from worker.pipeline.noo_rules import evaluate_noo_rubric
from worker.pipeline.docx_to_canonical import parse_docx_to_canonical


def test_noo_pipeline_smoke(tmp_path):
    p = tmp_path / "plan.docx"
    d = Document()
    d.add_paragraph("Тестовый план урока (CI)")
    d.add_paragraph("Цель урока: сформировать представление об алгоритме и научить применять его на практике.")
    d.add_paragraph("Задачи урока: 1) повторить понятия; 2) разобрать пример; 3) выполнить упражнение.")
    d.add_paragraph("Методы и приёмы: беседа, объяснение, практическая работа.")
    d.save(str(p))

    canonical = parse_docx_to_canonical(p)
    assert canonical["source"]["format"] == "docx"
    assert len(canonical.get("blocks") or []) >= 3

    extracted = extract_noo_from_canonical(canonical)
    assert isinstance(extracted, dict)

    rubric = load_rubric("noo_v1")
    results = evaluate_noo_rubric(rubric=rubric, canonical=canonical, extracted=extracted)

    assert results["max_score"] > 0
    assert 0 <= results["total_score"] <= results["max_score"]
    assert isinstance(results.get("criteria"), list)
    assert len(results["criteria"]) > 0
    assert isinstance(results.get("issues"), list)
