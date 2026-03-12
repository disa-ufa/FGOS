from __future__ import annotations

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from worker.pipeline.highlight_docx import highlight_docx_copy


def test_highlight_docx_copy_adds_legend_and_marks_requested_blocks(tmp_path):
    src = tmp_path / "plan.docx"
    dst = tmp_path / "plan_highlighted.docx"

    doc = Document()
    doc.add_paragraph("Тестовый конспект")          # p00000
    doc.add_paragraph("Цель урока")                 # p00001 -> RED
    doc.add_paragraph("Задачи урока")               # p00002 -> YELLOW
    doc.save(str(src))

    stats = highlight_docx_copy(
        src_path=str(src),
        dst_path=str(dst),
        severity_by_block_id={"p00001": 2, "p00002": 1},
        add_legend=True,
    )

    assert dst.exists()
    assert stats["blocks_marked"] == 2
    assert stats["paragraph_runs"] >= 2

    out = Document(str(dst))
    assert out.paragraphs[0].text.startswith("FGOS Helper: подсветка замечаний")

    target_goal = next(p for p in out.paragraphs if p.text == "Цель урока")
    target_tasks = next(p for p in out.paragraphs if p.text == "Задачи урока")

    assert any(run.font.highlight_color == WD_COLOR_INDEX.RED for run in target_goal.runs)
    assert any(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in target_tasks.runs)
