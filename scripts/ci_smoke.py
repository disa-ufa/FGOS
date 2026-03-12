from __future__ import annotations

import hashlib
import hmac
import os
import sys
import time
from io import BytesIO
from typing import Any, Dict

import requests
from docx import Document


def _sign(secret: str, method: str, path_qs: str) -> Dict[str, str]:
    ts = int(time.time())
    msg = f"{ts}.{method.upper()}.{path_qs}".encode("utf-8")
    sig = hmac.new(secret.encode("utf-8"), msg, hashlib.sha256).hexdigest()
    return {"X-Service-Timestamp": str(ts), "X-Service-Signature": sig}


def _req(base: str, secret: str, method: str, path_qs: str, **kwargs) -> requests.Response:
    headers = kwargs.pop("headers", {}) or {}
    headers.update(_sign(secret, method, path_qs))
    url = base.rstrip("/") + path_qs
    return requests.request(method, url, headers=headers, timeout=30, **kwargs)


def _make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Тестовый план урока (CI E2E)")
    doc.add_paragraph("Цель урока: сформировать представление об алгоритме и научить применять его на практике.")
    doc.add_paragraph("Задачи урока: 1) повторить понятия; 2) разобрать пример; 3) выполнить упражнение.")
    doc.add_paragraph("Методы и приёмы: беседа, объяснение, практическая работа.")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main() -> int:
    base = os.getenv("BASE_URL", "http://127.0.0.1:8000").strip()
    secret = os.getenv("SERVICE_SECRET", "").strip()
    if not secret:
        print("SERVICE_SECRET is required for CI smoke", file=sys.stderr)
        return 2

    chat_id = int(os.getenv("CI_CHAT_ID", "111"))
    # 1) Upload
    docx_bytes = _make_docx_bytes()
    files = {
        "file": (
            "ci_plan.docx",
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    data = {"telegram_user_id": "1", "chat_id": str(chat_id)}
    r = _req(base, secret, "POST", "/v1/documents", files=files, data=data)
    if r.status_code >= 400:
        print("Upload failed:", r.status_code, r.text, file=sys.stderr)
        return 3
    payload: Dict[str, Any] = r.json()
    job_id = payload["job_id"]
    print("Uploaded, job_id=", job_id)

    # 2) Poll job
    deadline = time.time() + 180
    status = None
    job = None
    while time.time() < deadline:
        jr = _req(base, secret, "GET", f"/v1/jobs/{job_id}?chat_id={chat_id}")
        if jr.status_code >= 400:
            print("Job get failed:", jr.status_code, jr.text, file=sys.stderr)
            return 4
        job = jr.json()
        status = job.get("status")
        if status in ("DONE", "FAILED"):
            break
        time.sleep(1)

    if status != "DONE":
        print("Job did not reach DONE. status=", status, "job=", job, file=sys.stderr)
        return 5

    # 3) Ensure pending-delivery exists
    pr = _req(base, secret, "GET", "/v1/bot/pending-deliveries?limit=20")
    if pr.status_code >= 400:
        print("pending-deliveries failed:", pr.status_code, pr.text, file=sys.stderr)
        return 6
    items = pr.json().get("items") or []
    assert any(str(it.get("job_id")) == str(job_id) and int(it.get("chat_id")) == chat_id for it in items), "delivery missing"

    # 4) Fetch bot job details (chat-scoped) and download artifacts
    br = _req(base, secret, "GET", f"/v1/bot/jobs/{job_id}?chat_id={chat_id}")
    if br.status_code >= 400:
        print("bot job failed:", br.status_code, br.text, file=sys.stderr)
        return 7
    bot_job = br.json()
    artifacts = bot_job.get("artifacts") or []
    if len(artifacts) < 3:
        print("Too few artifacts:", artifacts, file=sys.stderr)
        return 8

    for a in artifacts:
        aid = a["artifact_id"]
        dl_path = f"/v1/bot/jobs/{job_id}/artifacts/{aid}/download?chat_id={chat_id}"
        dr = _req(base, secret, "GET", dl_path)
        if dr.status_code >= 400:
            print("download failed:", dl_path, dr.status_code, dr.text, file=sys.stderr)
            return 9
        if len(dr.content) < 100:
            print("artifact too small:", a.get("filename"), len(dr.content), file=sys.stderr)
            return 10

    # 5) Ack
    ar = _req(base, secret, "POST", f"/v1/bot/deliveries/{job_id}/ack")
    if ar.status_code >= 400:
        print("ack failed:", ar.status_code, ar.text, file=sys.stderr)
        return 11

    print("CI smoke OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
